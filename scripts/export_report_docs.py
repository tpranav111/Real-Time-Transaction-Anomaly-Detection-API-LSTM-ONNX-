from __future__ import annotations

import argparse
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Export markdown technical report to DOCX and PDF."
    )
    parser.add_argument(
        "--input",
        default="docs/Fraud_Detection_System_Technical_Report.md",
        help="Input markdown file.",
    )
    parser.add_argument(
        "--docx",
        default="docs/Fraud_Detection_System_Technical_Report.docx",
        help="Output DOCX path.",
    )
    parser.add_argument(
        "--pdf",
        default="docs/Fraud_Detection_System_Technical_Report.pdf",
        help="Output PDF path.",
    )
    return parser.parse_args()


def export_docx(markdown_text: str, output_path: Path) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    in_code_block = False
    code_buffer: list[str] = []

    def flush_code() -> None:
        nonlocal code_buffer
        if not code_buffer:
            return
        p = doc.add_paragraph("\n".join(code_buffer))
        p.style = doc.styles["Normal"]
        for run in p.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        code_buffer = []

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                flush_code()
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            heading_text = stripped[level:].strip()
            doc.add_heading(heading_text, level=min(level, 4))
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        if (
            len(stripped) > 2
            and stripped[0].isdigit()
            and "." in stripped[:4]
            and stripped.split(".", 1)[1].strip()
        ):
            doc.add_paragraph(stripped.split(".", 1)[1].strip(), style="List Number")
            continue

        doc.add_paragraph(stripped)

    if in_code_block:
        flush_code()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def export_pdf(markdown_text: str, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=20,
        spaceBefore=12,
        spaceAfter=8,
    )
    h2 = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        spaceBefore=10,
        spaceAfter=6,
    )
    h3 = ParagraphStyle(
        "H3",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        spaceBefore=8,
        spaceAfter=4,
    )
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        spaceBefore=2,
        spaceAfter=4,
    )
    bullet = ParagraphStyle(
        "Bullet",
        parent=body,
        leftIndent=14,
    )
    code_style = ParagraphStyle(
        "Code",
        parent=body,
        fontName="Courier",
        fontSize=8.5,
        leading=11,
        leftIndent=10,
        backColor="#F3F3F3",
    )

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=LETTER,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36,
    )
    story = []

    in_code_block = False
    code_buffer: list[str] = []

    def flush_code() -> None:
        nonlocal code_buffer
        if not code_buffer:
            return
        story.append(Preformatted("\n".join(code_buffer), code_style))
        story.append(Spacer(1, 4))
        code_buffer = []

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                flush_code()
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        if not stripped:
            story.append(Spacer(1, 4))
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            heading_text = escape(stripped[level:].strip())
            style = h1 if level == 1 else h2 if level == 2 else h3
            story.append(Paragraph(heading_text, style))
            continue

        if stripped.startswith("- "):
            story.append(Paragraph(f"• {escape(stripped[2:].strip())}", bullet))
            continue

        if (
            len(stripped) > 2
            and stripped[0].isdigit()
            and "." in stripped[:4]
            and stripped.split(".", 1)[1].strip()
        ):
            body_text = escape(stripped)
            story.append(Paragraph(body_text, body))
            continue

        story.append(Paragraph(escape(stripped), body))

    if in_code_block:
        flush_code()

    doc.build(story)


def main() -> None:
    args = parse_args()
    input_path = Path(args.input)
    docx_path = Path(args.docx)
    pdf_path = Path(args.pdf)

    markdown_text = input_path.read_text(encoding="utf-8")
    export_docx(markdown_text, docx_path)
    export_pdf(markdown_text, pdf_path)

    print(f"docx={docx_path.resolve()}")
    print(f"pdf={pdf_path.resolve()}")


if __name__ == "__main__":
    main()
